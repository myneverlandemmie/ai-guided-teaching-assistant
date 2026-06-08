from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from openpyxl import load_workbook
from sqlalchemy import select

from app import main
from app.models.course import Course
from app.models.lesson import Lesson
from app.models.lesson_draft import LessonDraft
from app.services.ai.lesson_draft_service import build_chaoxing_catalog
from tests.support.course_plan_helpers import (
    _build_test_client,
    _create_course,
    _create_first_lesson,
    _create_reviewed_outline,
    anyio_backend,
    inline_threadpool_for_tests,
)


DOCX_SAMPLE_MARKDOWN = """# 学习导航

## 知识要点

- 传感器
- 数据采集

1. 阅读任务
2. 完成练习
(1) 查询所有数据
(2) 插入一条记录

本课学习目标：掌握数据表操作。
教师确认提示：请结合班级学情调整。
学生要做什么：
思考提示：
教师可调整点：
普通段落内容：**任务驱动**，请执行 `SELECT * FROM student;`。
请执行 `SELECT * FROM student;` 查看结果。
请运行 `print("hello")`。
请写出 `int main()` 的作用。
异常行：SELECT * FROM student;`

```sql
SELECT *
FROM student
WHERE id = 1;
```

```python
for i in range(3):
    print(i)
```

```c
int main() {
    return 0;
}
```

```markdown
# 一级标题
- 列表项
**加粗**
```

查看所有学生信息：
SELECT * FROM student;
插入一个新学生：
INSERT INTO student (id, name) VALUES (1, '张三');
"""


def _docx_part(docx_content: bytes, part_name: str) -> str:
    with ZipFile(BytesIO(docx_content)) as archive:
        return archive.read(part_name).decode("utf-8")


def _all_paragraphs(document: Document) -> list[object]:
    table_paragraphs = [
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    ]
    return list(document.paragraphs) + table_paragraphs


@pytest.mark.anyio
async def test_diagnostic_probe_exports_chaoxing_template(tmp_path: Path) -> None:
    client, session_factory = _build_test_client(tmp_path)
    course = _create_course(session_factory)
    try:
        await _create_first_lesson(client, session_factory, course)
        _create_reviewed_outline(session_factory)
        await client.post("/lessons/1/drafts/generate/diagnostic_probe", follow_redirects=False)
        with session_factory() as session:
            saved_course = session.get(Course, course.id)
            lesson = session.get(Lesson, 1)
            assert saved_course is not None
            assert lesson is not None
            saved_course.title = "数据库基础"
            lesson.lesson_code = "0406"
            lesson.title = "分组查询"
            session.commit()
            draft = session.scalar(select(LessonDraft).where(LessonDraft.draft_type == "diagnostic_probe"))
            assert draft is not None
            draft_id = draft.id

        export_response = await client.post(
            f"/lessons/1/drafts/{draft_id}/export-chaoxing",
            follow_redirects=False,
        )

        assert export_response.status_code == 303
        assert "chaoxing_file=" in export_response.headers["location"]
        export_files = list((tmp_path / "exports" / "chaoxing").glob("*.xlsx"))
        assert len(export_files) == 1
        workbook = load_workbook(export_files[0])
        worksheet = workbook["课程题库"]
        headers = [cell.value for cell in worksheet[1]]
        for header in ["目录", "题目类型", "大题题干", "正确答案", "答案解析", "难易度", "知识点", "标签", "选项数", "选项A", "选项B"]:
            assert header in headers
        rows = list(worksheet.iter_rows(min_row=2, values_only=True))
        assert len(rows) >= 5
        catalogs = {row[0] for row in rows}
        assert "/数据库基础/0406-分组查询" in catalogs
        assert all("智学导评" not in str(catalog) for catalog in catalogs)
        question_types = {row[1] for row in rows}
        assert {"单选题", "判断题", "填空题"}.issubset(question_types)
        judgment_row = next(row for row in rows if row[1] == "判断题")
        assert judgment_row[10] == 2
        assert judgment_row[11] == "正确"
        assert judgment_row[12] == "错误"

        page_response = await client.get(export_response.headers["location"])
        assert page_response.status_code == 200
        assert "下载学习通题库模板" in page_response.text
        assert "/exports/chaoxing/" in page_response.text
        assert "学习通 API" in page_response.text
    finally:
        await client.aclose()
        main.app.dependency_overrides.clear()


def test_chaoxing_catalog_falls_back_without_course_name() -> None:
    lesson = Lesson(
        id=1,
        course_id=1,
        planned_lesson_id=None,
        week="1",
        lesson_no="1",
        hours="2",
        lesson_code="0406",
        title="分组查询",
        content_summary="分组查询",
        status="draft",
    )

    assert build_chaoxing_catalog(lesson) == "/0406-分组查询"


@pytest.mark.anyio
async def test_guide_low_can_download_markdown(tmp_path: Path) -> None:
    client, session_factory = _build_test_client(tmp_path)
    course = _create_course(session_factory)
    try:
        await _create_first_lesson(client, session_factory, course)
        _create_reviewed_outline(session_factory)
        await client.post("/lessons/1/drafts/generate/guide_low", follow_redirects=False)
        with session_factory() as session:
            draft = session.scalar(select(LessonDraft).where(LessonDraft.draft_type == "guide_low"))
            assert draft is not None
            draft_id = draft.id
            expected_content = draft.content

        response = await client.get(f"/lessons/1/drafts/{draft_id}/download-md")

        assert response.status_code == 200
        assert "text/markdown" in response.headers["content-type"]
        assert ".md" in response.headers["content-disposition"]
        assert response.text == expected_content
        assert "学习导航" in response.text
        assert "rule-based" not in response.text
        assert "rule_based" not in response.text
        assert "mock" not in response.text
        markdown_file = tmp_path / "exports" / "guides" / "lesson_1_core_learning_guide.md"
        assert markdown_file.exists()
        assert markdown_file.read_text(encoding="utf-8") == expected_content
    finally:
        await client.aclose()
        main.app.dependency_overrides.clear()


@pytest.mark.anyio
async def test_guide_low_can_download_docx_with_basic_markdown(tmp_path: Path) -> None:
    client, session_factory = _build_test_client(tmp_path)
    course = _create_course(session_factory)
    try:
        await _create_first_lesson(client, session_factory, course)
        _create_reviewed_outline(session_factory)
        with session_factory() as session:
            saved_course = session.get(Course, course.id)
            lesson = session.get(Lesson, 1)
            assert saved_course is not None
            assert lesson is not None
            saved_course.title = "传感器应用基础"
            lesson.lesson_code = "0401"
            lesson.title = "光敏传感器数据采集"
            draft = LessonDraft(
                lesson_id=lesson.id,
                source_outline_id=None,
                draft_type="guide_low",
                title="0401-光敏传感器数据采集｜全班通用导学案",
                content=DOCX_SAMPLE_MARKDOWN,
                status="draft",
                generated_by="local-structured-draft",
            )
            session.add(draft)
            session.commit()
            draft_id = draft.id

        response = await client.get(f"/lessons/1/drafts/{draft_id}/download-docx")

        assert response.status_code == 200
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers["content-type"]
        assert ".docx" in response.headers["content-disposition"]
        document = Document(BytesIO(response.content))
        all_paragraphs = _all_paragraphs(document)
        document_text = "\n".join(paragraph.text for paragraph in all_paragraphs)
        header_text = "\n".join(paragraph.text for section in document.sections for paragraph in section.header.paragraphs)
        body_title = next(paragraph for paragraph in document.paragraphs if paragraph.style.name == "Title")
        assert "智学导评 V0.2" in header_text
        assert "AI 输出为教师草稿" in header_text
        assert body_title.text == "0401-光敏传感器数据采集｜全班通用导学案"
        assert "0401-光敏传感器数据采集｜0401-光敏传感器数据采集｜全班通用导学案" not in document_text
        assert not any(paragraph.style.name == "Title" and paragraph.text == "智学导评 V0.2" for paragraph in document.paragraphs)
        for expected_text in [
            "0401-光敏传感器数据采集｜全班通用导学案",
            "传感器应用基础",
            "光敏传感器数据采集",
            "学习导航",
            "知识要点",
            "传感器",
            "普通段落内容",
            "任务驱动",
            "SELECT * FROM student;",
            "SELECT *",
            "FROM student",
            "WHERE id = 1;",
            "for i in range(3):",
            "print(i)",
            "int main() {",
            "return 0;",
            "# 一级标题",
            "- 列表项",
            "**加粗**",
            "本课学习目标：掌握数据表操作。",
            "教师确认提示：请结合班级学情调整。",
            "学生要做什么：",
            "思考提示：",
            "教师可调整点：",
            "异常行：SELECT * FROM student;",
            "查询所有数据",
            "插入一条记录",
            "INSERT INTO student (id, name) VALUES (1, '张三');",
            'print("hello")',
            "int main()",
        ]:
            assert expected_text in document_text
        assert "**任务驱动**" not in document_text
        assert "`SELECT * FROM student;`" not in document_text
        assert '`print("hello")`' not in document_text
        assert "`int main()`" not in document_text
        assert "`" not in document_text
        assert "student;`" not in document_text
        assert "```" not in document_text
        assert "\nsql\n" not in f"\n{document_text}\n"
        assert "\npython\n" not in f"\n{document_text}\n"
        assert "\nc\n" not in f"\n{document_text}\n"
        assert "\nmarkdown\n" not in f"\n{document_text}\n"
        assert not any(paragraph.text.strip() in {"sql", "python", "c", "markdown"} for paragraph in all_paragraphs)
        styles_xml = _docx_part(response.content, "word/styles.xml")
        document_xml = _docx_part(response.content, "word/document.xml")
        assert 'w:eastAsia="宋体"' in styles_xml
        assert 'w:eastAsia="宋体"' in document_xml
        assert "w:keepNext" not in styles_xml
        assert "w:keepLines" not in styles_xml
        assert "w:keepNext" not in document_xml
        assert "w:keepLines" not in document_xml
        assert 'w:ascii="Consolas"' in document_xml
        assert 'w:eastAsia="Consolas"' in document_xml
        assert "<w:shd" in document_xml
        assert "<w:tbl>" in document_xml
        assert 'w:fill="EFEFEF"' in document_xml
        assert document.styles["Title"].font.bold is True
        assert document.styles["Heading 1"].font.bold is True
        assert document.styles["Heading 2"].font.size.pt < document.styles["Heading 1"].font.size.pt

        bold_runs = [run for paragraph in all_paragraphs for run in paragraph.runs if "任务驱动" in run.text]
        assert any(run.bold is True for run in bold_runs)
        inline_select_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "请执行 SELECT * FROM student; 查看结果。")
        assert any(run.text == "请执行 " and run.font.name == "Times New Roman" for run in inline_select_paragraph.runs)
        assert any(run.text == " 查看结果。" and run.font.name == "Times New Roman" for run in inline_select_paragraph.runs)
        inline_code_runs = [
            run
            for paragraph in document.paragraphs
            for run in paragraph.runs
            if run.text in {"SELECT * FROM student;", 'print("hello")', "int main()"}
        ]
        assert {run.text for run in inline_code_runs} == {"SELECT * FROM student;", 'print("hello")', "int main()"}
        assert all(run.font.name == "Consolas" for run in inline_code_runs)
        assert all('w:fill="EFEFEF"' in run._element.xml for run in inline_code_runs)
        code_block_paragraph = next(paragraph for paragraph in all_paragraphs if "WHERE id = 1;" in paragraph.text)
        assert "SELECT *" in code_block_paragraph.text
        assert "FROM student" in code_block_paragraph.text
        assert all(run.font.name == "Consolas" for run in code_block_paragraph.runs if run.text.strip())
        table_texts = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        assert any("SELECT *\nFROM student\nWHERE id = 1;" in text for text in table_texts)
        assert any("for i in range(3):\n    print(i)" in text for text in table_texts)
        assert any("int main() {\n    return 0;\n}" in text for text in table_texts)
        assert any("# 一级标题\n- 列表项\n**加粗**" in text for text in table_texts)
        assert any(text == "SELECT * FROM student;" for text in table_texts)
        assert any("INSERT INTO student (id, name) VALUES (1, '张三');" in text for text in table_texts)
        assert all("查看所有学生信息：" not in text for text in table_texts)
        assert all("插入一个新学生：" not in text for text in table_texts)

        natural_paragraphs = [
            next(paragraph for paragraph in document.paragraphs if paragraph.text == text)
            for text in [
                "本课学习目标：掌握数据表操作。",
                "教师确认提示：请结合班级学情调整。",
                "查看所有学生信息：",
                "插入一个新学生：",
            ]
        ]
        label_paragraphs = [
            next(paragraph for paragraph in document.paragraphs if paragraph.text == text)
            for text in ["学生要做什么：", "思考提示：", "教师可调整点："]
        ]
        assert {paragraph.style.name for paragraph in natural_paragraphs + label_paragraphs} == {"Normal"}
        assert all(run.font.name == "Times New Roman" for paragraph in natural_paragraphs for run in paragraph.runs if run.text)
        assert all(any(run.bold is True for run in paragraph.runs if run.text) for paragraph in label_paragraphs)
        bullet_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text in {"传感器", "数据采集"}]
        number_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text in {"阅读任务", "完成练习", "查询所有数据", "插入一条记录"}
        ]
        assert {paragraph.style.name for paragraph in bullet_paragraphs} == {"List Bullet"}
        assert {paragraph.style.name for paragraph in number_paragraphs} == {"List Number"}
        docx_files = list((tmp_path / "exports" / "guides").glob("*.docx"))
        assert len(docx_files) == 1
        assert docx_files[0].name == "lesson_1_0401_core_learning_guide.docx"
    finally:
        await client.aclose()
        main.app.dependency_overrides.clear()


@pytest.mark.anyio
async def test_docx_download_empty_draft_shows_friendly_message(tmp_path: Path) -> None:
    client, session_factory = _build_test_client(tmp_path)
    course = _create_course(session_factory)
    try:
        await _create_first_lesson(client, session_factory, course)
        with session_factory() as session:
            draft = LessonDraft(
                lesson_id=1,
                source_outline_id=None,
                draft_type="guide_low",
                title="空草稿",
                content="   ",
                status="draft",
                generated_by="local-structured-draft",
            )
            session.add(draft)
            session.commit()
            draft_id = draft.id

        response = await client.get(f"/lessons/1/drafts/{draft_id}/download-docx")

        assert response.status_code == 400
        assert "DOCX 文件生成失败，请重新生成或稍后再试。" in response.text
        assert "Traceback" not in response.text
        assert "/home/" not in response.text
    finally:
        await client.aclose()
        main.app.dependency_overrides.clear()

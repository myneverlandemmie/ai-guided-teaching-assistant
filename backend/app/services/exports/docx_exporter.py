"""基础 DOCX 导出工具。"""

from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.models.lesson import Lesson
from app.models.lesson_draft import LessonDraft


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOCX_EXPORT_ERROR_MESSAGE = "DOCX 文件生成失败，请重新生成或稍后再试。"
NORMAL_EAST_ASIA_FONT = "宋体"
NORMAL_LATIN_FONT = "Times New Roman"
CODE_FONT = "Consolas"
CODE_BLOCK_FILL = "EFEFEF"

_HEADING_PATTERN = re.compile(r"^(#{1,3})\s+(.+)$")
_ORDERED_LIST_PATTERN = re.compile(r"^(?:\d+[\.)]|\(\d+\))\s+(.+)$")
_UNORDERED_LIST_PATTERN = re.compile(r"^[-*]\s+(.+)$")
_SQL_LINE_PATTERN = re.compile(r"^(SELECT|INSERT|UPDATE|DELETE|TRUNCATE|CREATE|ALTER|DROP)\b", re.IGNORECASE)
_TASK_HEADING_PATTERN = re.compile(r"^任务\s*\d+\s*[:：].+")
_CHINESE_LABEL_PATTERN = re.compile(r"^[\u4e00-\u9fffA-Za-z0-9（）()、/／\s]+[:：]\s*$")


class DocxExportError(Exception):
    """DOCX 生成失败。"""


def build_lesson_draft_docx(lesson: Lesson, draft: LessonDraft) -> bytes:
    """将单个课次草稿转换为基础 DOCX 字节。"""

    content = (draft.content or "").strip()
    if not content:
        raise DocxExportError(DOCX_EXPORT_ERROR_MESSAGE)

    document = Document()
    _configure_document_styles(document)
    _set_header_text(document, "智学导评 V0.2｜AI 输出为教师草稿，需教师审阅确认后使用")
    lesson_label = _join_non_empty([lesson.lesson_code, lesson.title], "-")
    document_title = _build_document_title(lesson, draft)
    document.core_properties.title = document_title
    _add_markdown_paragraph(document, document_title, style="Title")

    _add_markdown_paragraph(document, "课次信息", style="Heading 1")
    course = lesson.course
    if course is not None:
        _add_markdown_paragraph(document, f"课程：{course.title}")
        if course.semester:
            _add_markdown_paragraph(document, f"学期：{course.semester}")
    _add_markdown_paragraph(document, f"课次：{lesson_label or f'课次 {lesson.id}'}")
    _add_markdown_paragraph(document, f"周次：{lesson.week or '-'}")
    _add_markdown_paragraph(document, f"课时：{lesson.hours or '-'}")

    _add_markdown_paragraph(document, draft.title or "草稿内容", style="Heading 1")
    _append_markdown_content(document, content)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _append_markdown_content(document: Document, content: str) -> None:
    """将轻量 Markdown 结构写入 DOCX。"""

    in_code_block = False
    code_lines: list[str] = []
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if line.startswith("```"):
            if in_code_block:
                _add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(raw_line)
            continue

        if not line:
            continue

        heading_match = _HEADING_PATTERN.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            _add_markdown_paragraph(document, heading_match.group(2).strip(), _markdown_heading_style(level))
            continue

        unordered_match = _UNORDERED_LIST_PATTERN.match(raw_line)
        if unordered_match:
            _add_markdown_paragraph(document, unordered_match.group(1).strip(), "List Bullet")
            continue

        ordered_match = _ORDERED_LIST_PATTERN.match(raw_line)
        if ordered_match:
            _add_markdown_paragraph(document, ordered_match.group(1).strip(), "List Number")
            continue

        if _SQL_LINE_PATTERN.match(line):
            _add_code_block(document, [line])
            continue

        if _TASK_HEADING_PATTERN.match(line):
            _add_markdown_paragraph(document, line, "Heading 2")
            continue

        if _CHINESE_LABEL_PATTERN.match(line):
            _add_markdown_paragraph(document, line, force_bold=True)
            continue

        _add_markdown_paragraph(document, line)

    if in_code_block:
        _add_code_block(document, code_lines)


def _configure_document_styles(document: Document) -> None:
    """设置 DOCX 的基础中文字体和标题层级。"""

    _set_style_font(document, "Normal", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 11, bold=False)
    _set_style_font(document, "Title", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 22, bold=True, color="1F3B35")
    _set_style_font(document, "Heading 1", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 16, bold=True, color="1F3B35")
    _set_style_font(document, "Heading 2", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 14, bold=True, color="234942")
    _set_style_font(document, "Heading 3", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 12, bold=True, color="2F5B53")
    _set_style_font(document, "List Bullet", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 11, bold=False)
    _set_style_font(document, "List Number", NORMAL_EAST_ASIA_FONT, NORMAL_LATIN_FONT, 11, bold=False)

    _set_paragraph_spacing(document, "Normal", after=4)
    _set_paragraph_spacing(document, "Title", before=0, after=12)
    _set_paragraph_spacing(document, "Heading 1", before=14, after=8)
    _set_paragraph_spacing(document, "Heading 2", before=10, after=6)
    _set_paragraph_spacing(document, "Heading 3", before=8, after=4)
    _set_paragraph_spacing(document, "List Bullet", after=2)
    _set_paragraph_spacing(document, "List Number", after=2)
    _remove_keep_properties(document)


def _set_header_text(document: Document, text: str) -> None:
    header = document.sections[0].header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(4)
    _add_text_run(paragraph, text, size_pt=9)


def _markdown_heading_style(level: int) -> str:
    return "Heading 2" if level <= 2 else "Heading 3"


def _add_markdown_paragraph(document: Document, text: str, style: str | None = None, *, force_bold: bool = False) -> None:
    try:
        paragraph = document.add_paragraph(style=style)
    except KeyError:
        paragraph = document.add_paragraph()
    if force_bold:
        _add_text_run(paragraph, text, bold=True)
        return
    _add_inline_markdown_runs(paragraph, text)


def _add_inline_markdown_runs(paragraph: object, text: str) -> None:
    index = 0
    while index < len(text):
        if text.startswith("**", index):
            end_index = text.find("**", index + 2)
            if end_index == -1:
                _add_text_run(paragraph, text[index:])
                break
            run = paragraph.add_run(text[index + 2 : end_index])
            _set_run_font(run, NORMAL_LATIN_FONT, NORMAL_EAST_ASIA_FONT, bold=True)
            index = end_index + 2
            continue

        if text[index] == "`":
            end_index = text.find("`", index + 1)
            if end_index == -1:
                if index + 1 < len(text):
                    _add_text_run(paragraph, text[index:])
                break
            _add_code_run(paragraph, text[index + 1 : end_index], shade=True)
            index = end_index + 1
            continue

        next_bold = text.find("**", index)
        next_code = text.find("`", index)
        next_candidates = [position for position in [next_bold, next_code] if position != -1]
        next_index = min(next_candidates) if next_candidates else len(text)
        _add_text_run(paragraph, text[index:next_index])
        index = next_index


def _add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.allow_autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, CODE_BLOCK_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)

    if not lines:
        _add_code_run(paragraph, "")
        return

    for index, line in enumerate(lines):
        if index > 0:
            paragraph.add_run().add_break()
        _add_code_run(paragraph, line)


def _add_code_run(paragraph: object, text: str, *, shade: bool = False) -> None:
    run = paragraph.add_run(text)
    _set_run_font(run, CODE_FONT, CODE_FONT, 10)
    if shade:
        _set_run_shading(run, CODE_BLOCK_FILL)


def _add_text_run(paragraph: object, text: str, *, bold: bool = False, size_pt: int | None = None) -> None:
    run = paragraph.add_run(text)
    _set_run_font(run, NORMAL_LATIN_FONT, NORMAL_EAST_ASIA_FONT, size_pt=size_pt, bold=bold)


def _set_style_font(
    document: Document,
    style_name: str,
    east_asia_font: str,
    latin_font: str,
    size_pt: int,
    *,
    bold: bool,
    color: str | None = None,
) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    style.font.name = latin_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    _set_rpr_fonts(style.element.get_or_add_rPr(), latin_font, east_asia_font)


def _set_paragraph_spacing(document: Document, style_name: str, *, before: int | None = None, after: int | None = None) -> None:
    try:
        paragraph_format = document.styles[style_name].paragraph_format
    except KeyError:
        return
    if before is not None:
        paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = 1.15


def _remove_keep_properties(document: Document) -> None:
    for style in document.styles:
        paragraph_properties = getattr(style.element, "pPr", None)
        if paragraph_properties is None:
            continue
        _remove_child(paragraph_properties, "w:keepNext")
        _remove_child(paragraph_properties, "w:keepLines")


def _remove_child(parent: object, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def _set_run_font(
    run: object,
    latin_font: str,
    east_asia_font: str,
    size_pt: int | None = None,
    *,
    bold: bool | None = None,
) -> None:
    run.font.name = latin_font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    _set_rpr_fonts(run._element.get_or_add_rPr(), latin_font, east_asia_font)


def _set_rpr_fonts(rpr: object, latin_font: str, east_asia_font: str) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def _set_cell_shading(cell: object, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def _set_run_shading(run: object, fill: str) -> None:
    run_properties = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    run_properties.append(shading)


def _join_non_empty(parts: list[str | None], separator: str) -> str:
    return separator.join(part.strip() for part in parts if part and part.strip())


def _build_document_title(lesson: Lesson, draft: LessonDraft) -> str:
    draft_title = (draft.title or "教师草稿").strip()
    lesson_code = (lesson.lesson_code or "").strip()
    lesson_title = (lesson.title or "").strip()
    lesson_label = _join_non_empty([lesson_code, lesson_title], "-")
    if _title_contains_lesson_identity(draft_title, lesson_code, lesson_title, lesson_label):
        return draft_title
    return _join_non_empty([lesson_label, draft_title], "｜") or "教师草稿"


def _title_contains_lesson_identity(title: str, lesson_code: str, lesson_title: str, lesson_label: str) -> bool:
    title_without_spaces = title.replace(" ", "")
    for value in [lesson_label, lesson_code, lesson_title]:
        normalized = value.replace(" ", "")
        if normalized and normalized in title_without_spaces:
            return True
    return False
